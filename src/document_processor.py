import re

from docx import Document

from detector import detect_regex_pii
from name_detector import detect_ner_pii
from person_name_detector import (
    detect_context_persons,
)


def remove_overlaps(findings):

    confidence_order = {
        "HIGH": 0,
        "MEDIUM": 1,
        "LOW": 2,
    }

    # Higher-confidence detections first.
    findings.sort(
        key=lambda item: (
            item["start"],
            confidence_order.get(
                item.get("confidence"),
                99
            ),
            -(item["end"] - item["start"])
        )
    )

    result = []

    occupied_until = -1

    for finding in findings:

        if finding["start"] >= occupied_until:

            result.append(finding)

            occupied_until = finding["end"]

    return result


def process_text(text, anonymizer):

    if not text.strip():
        return text, []

    findings = []

    # ---------------------------------------------------------
    # 1. Regex / structured PII
    # ---------------------------------------------------------

    findings.extend(
        detect_regex_pii(text)
    )

    # ---------------------------------------------------------
    # 2. spaCy NER
    # ---------------------------------------------------------

    findings.extend(
        detect_ner_pii(text)
    )

    # ---------------------------------------------------------
    # 3. Context-based PERSON detection
    # ---------------------------------------------------------

    findings.extend(
        detect_context_persons(text)
    )

    # ---------------------------------------------------------
    # Remove already-generated fake values
    # ---------------------------------------------------------

    findings = [
        finding
        for finding in findings
        if not anonymizer.is_generated_value(
            finding["value"]
        )
    ]

    # ---------------------------------------------------------
    # Remove overlapping detections
    # ---------------------------------------------------------

    findings = remove_overlaps(
        findings
    )

    # ---------------------------------------------------------
    # Keep a copy for the document-level audit.
    #
    # We need the original findings before replacing text.
    # ---------------------------------------------------------

    detected_findings = list(
        findings
    )

    # ---------------------------------------------------------
    # Replace from RIGHT to LEFT
    #
    # This is important because replacing text from
    # left to right changes character positions.
    # ---------------------------------------------------------

    findings.sort(
        key=lambda item: item["start"],
        reverse=True
    )

    result = text

    for finding in findings:

        original = finding["value"]

        replacement = anonymizer.get_replacement(
            finding["type"],
            original
        )

        result = (
            result[:finding["start"]]
            + replacement
            + result[finding["end"]:]
        )

    # ---------------------------------------------------------
    # Return BOTH:
    #
    # 1. processed text
    # 2. what was detected
    # ---------------------------------------------------------

    return result, detected_findings

def is_generated_value(self, value):

    normalized_value = self.normalize(
        value
    )

    for generated in self.generated_values:

        if self.normalize(
            generated
        ) == normalized_value:

            return True

    return False


def print_detection_summary(all_findings):

    print()
    print("=" * 60)
    print("FINAL DETECTION SUMMARY")
    print("=" * 60)

    categories = [
        "PERSON",
        "ORGANIZATION",
        "EMAIL",
        "PHONE",
        "ADDRESS",
        "SSN",
        "CREDIT_CARD",
        "DATE",
        "IP_ADDRESS",
    ]

    # ---------------------------------------------------------
    # Category totals
    # ---------------------------------------------------------

    for category in categories:

        category_findings = [
            finding
            for finding in all_findings
            if finding.get("type") == category
        ]

        unique_values = {}

        for finding in category_findings:

            value = re.sub(
                r"\s+",
                " ",
                finding.get(
                    "value",
                    ""
                ).strip()
            )

            if not value:
                continue

            key = value.casefold()

            if key not in unique_values:

                unique_values[key] = value

        print(
            f"{category:<15}: "
            f"{len(unique_values)}"
        )

    # ---------------------------------------------------------
    # PERSON details
    # ---------------------------------------------------------

    person_findings = [
        finding
        for finding in all_findings
        if finding.get("type") == "PERSON"
    ]

    unique_persons = {}

    for finding in person_findings:

        value = re.sub(
            r"\s+",
            " ",
            finding.get(
                "value",
                ""
            ).strip()
        )

        if not value:
            continue

        key = value.casefold()

        if key not in unique_persons:

            unique_persons[key] = {
                "value": value,
                "source": finding.get(
                    "source",
                    "unknown"
                ),
                "confidence": finding.get(
                    "confidence",
                    "unknown"
                ),
            }

    print()
    print("PERSON DETAILS")
    print("-" * 60)

    print(
        f"Unique persons detected: "
        f"{len(unique_persons)}"
    )

    # ---------------------------------------------------------
    # Detection source counts
    # ---------------------------------------------------------

    source_counts = {}

    for person in unique_persons.values():

        source = person["source"]

        source_counts[source] = (
            source_counts.get(
                source,
                0
            ) + 1
        )

    if source_counts:

        print()
        print("Detection sources:")

        for source, count in sorted(
            source_counts.items()
        ):

            print(
                f"  {source:<15}: "
                f"{count}"
            )

    # ---------------------------------------------------------
    # Show detected names
    #
    # Limit this to 30 so the terminal doesn't become huge.
    # ---------------------------------------------------------

    print()
    print(
        "Detected PERSON names "
        "(first 30):"
    )

    for person in list(
        unique_persons.values()
    )[:30]:

        print(
            f"  - {person['value']} "
            f"[{person['source']}]"
        )

    if len(unique_persons) > 30:

        print(
            f"  ... and "
            f"{len(unique_persons) - 30} "
            f"more"
        )

    print("=" * 60)
    print()


def redact_docx(
    input_path,
    output_path,
    anonymizer
):

    document = Document(
        input_path
    )

    # ---------------------------------------------------------
    # Store ALL detections from the entire document.
    # ---------------------------------------------------------

    all_findings = []

    # ---------------------------------------------------------
    # Normal paragraphs
    # ---------------------------------------------------------

    for paragraph in document.paragraphs:

        processed_text, findings = process_text(
            paragraph.text,
            anonymizer
        )

        paragraph.text = processed_text

        all_findings.extend(
            findings
        )

    # ---------------------------------------------------------
    # Tables
    # ---------------------------------------------------------

    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                processed_text, findings = process_text(
                    cell.text,
                    anonymizer
                )

                cell.text = processed_text

                all_findings.extend(
                    findings
                )

    # ---------------------------------------------------------
    # Save document
    # ---------------------------------------------------------

    document.save(
        output_path
    )

    # ---------------------------------------------------------
    # Print ONE summary for the whole document.
    # ---------------------------------------------------------

    print_detection_summary(
        all_findings
    )